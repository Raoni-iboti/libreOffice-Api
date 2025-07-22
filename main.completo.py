from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
import os
import shutil
import re
from docx import Document
from typing import Dict, List
from pydantic import BaseModel

app = FastAPI(title="API de Substituição com Marcadores - CORRIGIDA", version="2.1.0")
UPLOAD_PATH = "temp"
INPUT_FILE = os.path.join(UPLOAD_PATH, "original.docx")
OUTPUT_FILE = os.path.join(UPLOAD_PATH, "editado.docx")

os.makedirs(UPLOAD_PATH, exist_ok=True)

class SubstituicoesModel(BaseModel):
    substituicoes: Dict[str, str]

def encontrar_marcadores(doc: Document) -> List[str]:
    """
    Encontra todos os marcadores no documento.
    Versão corrigida que evita duplicatas.
    """
    marcadores = set()
    
    # Padrões de marcadores suportados (corrigidos)
    padroes = [
        r'\{\{([^}]+)\}\}',  # {{NOME}}
        r'\[([^\]]+)\]',     # [NOME]  
        r'<<([^>]+)>>',      # <<NOME>>
        r'\{([^}]+)\}'       # {NOME} - apenas chaves simples
    ]
    
    def extrair_marcadores_texto(texto: str):
        for padrao in padroes:
            matches = re.findall(padrao, texto)
            for match in matches:
                # Limpar e validar marcador
                marcador = match.strip()
                if marcador and not marcador.startswith('{') and not marcador.startswith('['):
                    marcadores.add(marcador)
    
    # Buscar em parágrafos
    for paragraph in doc.paragraphs:
        extrair_marcadores_texto(paragraph.text)
    
    # Buscar em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extrair_marcadores_texto(cell.text)
    
    # Buscar em cabeçalhos e rodapés
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                extrair_marcadores_texto(paragraph.text)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                extrair_marcadores_texto(paragraph.text)
    
    return sorted(list(marcadores))

def substituir_marcadores(doc: Document, substituicoes: Dict[str, str]) -> int:
    """
    Substitui marcadores no documento pelos valores fornecidos.
    Versão corrigida.
    """
    total_substituicoes = 0
    
    def substituir_em_texto(texto: str) -> tuple[str, int]:
        texto_modificado = texto
        subs_locais = 0
        
        for marcador, valor in substituicoes.items():
            # Padrões para substituição
            padroes_busca = [
                f'{{{{{marcador}}}}}',  # {{MARCADOR}}
                f'[{marcador}]',        # [MARCADOR]
                f'<<{marcador}>>',      # <<MARCADOR>>
                f'{{{marcador}}}'       # {MARCADOR}
            ]
            
            for padrao in padroes_busca:
                if padrao in texto_modificado:
                    texto_modificado = texto_modificado.replace(padrao, valor)
                    subs_locais += 1
        
        return texto_modificado, subs_locais
    
    # Substituir em parágrafos
    for paragraph in doc.paragraphs:
        texto_novo, subs = substituir_em_texto(paragraph.text)
        if subs > 0:
            paragraph.text = texto_novo
            total_substituicoes += subs
    
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto_novo, subs = substituir_em_texto(cell.text)
                if subs > 0:
                    cell.text = texto_novo
                    total_substituicoes += subs
    
    # Substituir em cabeçalhos e rodapés
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                texto_novo, subs = substituir_em_texto(paragraph.text)
                if subs > 0:
                    paragraph.text = texto_novo
                    total_substituicoes += subs
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                texto_novo, subs = substituir_em_texto(paragraph.text)
                if subs > 0:
                    paragraph.text = texto_novo
                    total_substituicoes += subs
    
    return total_substituicoes

@app.post("/abrir")
async def abrir_docx(file: UploadFile = File(...)):
    """
    Faz upload de um arquivo DOCX e identifica marcadores automaticamente.
    """
    try:
        with open(INPUT_FILE, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        shutil.copy(INPUT_FILE, OUTPUT_FILE)
        
        # Identificar marcadores no documento
        doc = Document(INPUT_FILE)
        marcadores = encontrar_marcadores(doc)
        
        return {
            "status": "ok", 
            "message": "Arquivo salvo",
            "marcadores_encontrados": marcadores,
            "total_marcadores": len(marcadores)
        }
    except Exception as e:
        return {"status": "error", "message": f"Erro ao salvar arquivo: {str(e)}"}

@app.get("/marcadores")
async def listar_marcadores():
    """
    Lista todos os marcadores encontrados no documento carregado.
    """
    try:
        if not os.path.exists(INPUT_FILE):
            raise HTTPException(status_code=400, detail="Nenhum arquivo carregado. Use /abrir primeiro.")
        
        doc = Document(INPUT_FILE)
        marcadores = encontrar_marcadores(doc)
        
        return {
            "status": "ok",
            "marcadores": marcadores,
            "total": len(marcadores),
            "formatos_suportados": [
                "{{NOME}} - Chaves duplas (recomendado)",
                "[NOME] - Colchetes",
                "<<NOME>> - Sinais de menor/maior",
                "{NOME} - Chaves simples"
            ]
        }
    except Exception as e:
        return {"status": "error", "message": f"Erro ao listar marcadores: {str(e)}"}

@app.post("/substituir")
async def substituir_texto(buscar: str = Form(...), por: str = Form(...)):
    """
    Substitui texto específico no documento (modo manual).
    """
    try:
        if not os.path.exists(OUTPUT_FILE):
            raise HTTPException(status_code=400, detail="Nenhum arquivo carregado. Use /abrir primeiro.")
        
        doc = Document(OUTPUT_FILE)
        substituicoes_realizadas = 0
        
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            if buscar in paragraph.text:
                paragraph.text = paragraph.text.replace(buscar, por)
                substituicoes_realizadas += 1
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if buscar in cell.text:
                        cell.text = cell.text.replace(buscar, por)
                        substituicoes_realizadas += 1
        
        doc.save(OUTPUT_FILE)
        
        return {
            "status": "ok", 
            "message": "Texto substituído",
            "substituicoes": substituicoes_realizadas
        }
    except Exception as e:
        return {"status": "error", "message": f"Erro ao substituir texto: {str(e)}"}

@app.post("/substituir-marcadores")
async def substituir_marcadores_endpoint(dados: SubstituicoesModel):
    """
    Substitui marcadores usando um modelo Pydantic.
    
    Exemplo de uso:
    {
        "substituicoes": {
            "NOME": "João Silva",
            "DATA": "22/07/2025",
            "EMPRESA": "Minha Empresa Ltda"
        }
    }
    """
    try:
        if not os.path.exists(OUTPUT_FILE):
            raise HTTPException(status_code=400, detail="Nenhum arquivo carregado. Use /abrir primeiro.")
        
        doc = Document(OUTPUT_FILE)
        total_substituicoes = substituir_marcadores(doc, dados.substituicoes)
        doc.save(OUTPUT_FILE)
        
        return {
            "status": "ok",
            "message": "Marcadores substituídos com sucesso",
            "substituicoes_realizadas": total_substituicoes,
            "marcadores_processados": list(dados.substituicoes.keys())
        }
    except Exception as e:
        return {"status": "error", "message": f"Erro ao substituir marcadores: {str(e)}"}

@app.get("/baixar")
async def baixar():
    """
    Baixa o arquivo DOCX editado.
    """
    try:
        if not os.path.exists(OUTPUT_FILE):
            raise HTTPException(status_code=404, detail="Arquivo não encontrado")
        
        return FileResponse(
            OUTPUT_FILE, 
            filename="editado.docx", 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao baixar arquivo: {str(e)}")

@app.get("/")
async def root():
    """
    Página inicial com instruções de uso.
    """
    return {
        "message": "API de Substituição com Marcadores - VERSÃO CORRIGIDA",
        "versao": "2.1.0",
        "endpoints": {
            "POST /abrir": "Upload do arquivo DOCX (identifica marcadores automaticamente)",
            "GET /marcadores": "Lista marcadores encontrados no documento",
            "POST /substituir": "Substituição manual de texto",
            "POST /substituir-marcadores": "Substituição automática usando JSON",
            "GET /baixar": "Download do arquivo editado"
        },
        "formatos_marcadores": [
            "{{NOME}} - Recomendado",
            "[NOME] - Alternativo", 
            "<<NOME>> - Alternativo",
            "{NOME} - Simples"
        ],
        "exemplo_uso": {
            "substituicoes": {
                "NOME": "João Silva",
                "DATA": "22/07/2025",
                "EMPRESA": "Minha Empresa"
            }
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

