from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
import os
import shutil
from docx import Document

app = FastAPI()
UPLOAD_PATH = "temp"
INPUT_FILE = os.path.join(UPLOAD_PATH, "original.docx")
OUTPUT_FILE = os.path.join(UPLOAD_PATH, "editado.docx")

os.makedirs(UPLOAD_PATH, exist_ok=True)

@app.post("/abrir")
async def abrir_docx(file: UploadFile = File(...)):
    """
    Endpoint para fazer upload de um arquivo DOCX.
    Mantém a mesma funcionalidade do código original.
    """
    try:
        with open(INPUT_FILE, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        shutil.copy(INPUT_FILE, OUTPUT_FILE)
        return {"status": "ok", "message": "Arquivo salvo"}
    except Exception as e:
        return {"status": "error", "message": f"Erro ao salvar arquivo: {str(e)}"}

@app.post("/substituir")
async def substituir_texto(buscar: str = Form(...), por: str = Form(...)):
    """
    Endpoint para substituir texto no arquivo DOCX.
    Versão corrigida usando python-docx em vez do comando LibreOffice problemático.
    """
    try:
        # Verificar se o arquivo existe
        if not os.path.exists(OUTPUT_FILE):
            return {"status": "error", "message": "Nenhum arquivo carregado. Use /abrir primeiro."}
        
        # Carregar o documento
        doc = Document(OUTPUT_FILE)
        
        # Substituir texto em todos os parágrafos
        for paragraph in doc.paragraphs:
            if buscar in paragraph.text:
                paragraph.text = paragraph.text.replace(buscar, por)
        
        # Substituir texto em todas as tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if buscar in cell.text:
                        cell.text = cell.text.replace(buscar, por)
        
        # Substituir texto em cabeçalhos
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if buscar in paragraph.text:
                        paragraph.text = paragraph.text.replace(buscar, por)
        
        # Substituir texto em rodapés
        for section in doc.sections:
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if buscar in paragraph.text:
                        paragraph.text = paragraph.text.replace(buscar, por)
        
        # Salvar o documento modificado
        doc.save(OUTPUT_FILE)
        
        return {"status": "ok", "message": "Texto substituído"}
    
    except Exception as e:
        return {"status": "error", "message": f"Erro ao substituir texto: {str(e)}"}

@app.get("/baixar")
async def baixar():
    """
    Endpoint para baixar o arquivo DOCX editado.
    Mantém a mesma funcionalidade do código original.
    """
    try:
        if not os.path.exists(OUTPUT_FILE):
            raise HTTPException(status_code=404, detail="Arquivo não encontrado")
        
        return FileResponse(
            OUTPUT_FILE, 
            filename="editado.docx", 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao baixar arquivo: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

